from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/tarunydv/Desktop/Code/AstrologyApp")
OUT = ROOT / "artifacts/parashari-video/Parashari_Desk_Video_Production_Script.docx"
SHOTS = ROOT / "artifacts/parashari-desk-guide/screenshots"

ROSE = "A5123B"
DARK_ROSE = "7F0D2D"
INK = "302724"
MUTED = "766761"
IVORY = "FFF9F6"
SAND = "F2E9E2"
PALE_ROSE = "F9E8EE"
GREEN = "397050"
WHITE = "FFFFFF"
LINE = "DCCFC6"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, props in edges.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in props.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(ROSE)
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_callout(doc, title, text, fill=PALE_ROSE, accent=ROSE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 9120)
    shade(cell, fill)
    set_cell_margins(cell, 120, 170, 120, 170)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": accent},
        top={"val": "single", "sz": "4", "color": LINE},
        bottom={"val": "single", "sz": "4", "color": LINE},
        right={"val": "single", "sz": "4", "color": LINE},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    for r2 in p2.runs:
        r2.font.size = Pt(9.4)
        r2.font.color.rgb = RGBColor.from_string(INK)
    return table


def add_scene_header(doc, number, title, timecode, duration, objective, start_state):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"SCENE {number:02d}  /  ")
    r.font.color.rgb = RGBColor.from_string(ROSE)
    r2 = p.add_run(title)
    r2.font.color.rgb = RGBColor.from_string(INK)
    add_label_value(doc, "Time", f"{timecode}  ·  {duration}")
    add_label_value(doc, "Goal", objective)
    add_label_value(doc, "Starting state", start_state)


def add_cue_table(doc, rows):
    p = doc.add_paragraph("Synchronized cue sheet", style="Heading 2")
    p.paragraph_format.space_before = Pt(7)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1050, 4100, 3970]
    headers = ["RELATIVE TIME", "SCREEN / CURSOR / EDIT", "AMAN NARRATION CUE"]
    for i, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
        set_cell_width(cell, width)
        shade(cell, ROSE)
        set_cell_margins(cell, 85, 110, 85, 110)
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        run = p0.add_run(text)
        run.bold = True
        run.font.size = Pt(7.6)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, (t, action, cue) in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, (cell, width, value) in enumerate(zip(cells, widths, (t, action, cue))):
            set_cell_width(cell, width)
            set_cell_margins(cell, 85, 110, 85, 110)
            shade(cell, IVORY if idx % 2 == 0 else WHITE)
            p0 = cell.paragraphs[0]
            p0.paragraph_format.space_after = Pt(0)
            p0.paragraph_format.line_spacing = 1.08
            run = p0.add_run(value)
            run.font.size = Pt(8.2 if i else 7.8)
            run.font.color.rgb = RGBColor.from_string(INK if i else ROSE)
            if i == 0:
                run.bold = True
            set_cell_border(
                cell,
                top={"val": "single", "sz": "3", "color": LINE},
                bottom={"val": "single", "sz": "3", "color": LINE},
                left={"val": "single", "sz": "3", "color": LINE},
                right={"val": "single", "sz": "3", "color": LINE},
            )
    return table


def add_narration(doc, text):
    add_callout(doc, "Exact narration", text, fill="FFF4F7", accent=ROSE)


def add_editor_notes(doc, notes):
    p = doc.add_paragraph("Editor notes", style="Heading 2")
    p.paragraph_format.space_before = Pt(7)
    for note in notes:
        q = doc.add_paragraph(style="List Bullet")
        q.paragraph_format.space_after = Pt(2)
        run = q.add_run(note)
        run.font.size = Pt(8.8)


def add_image(doc, filename, caption, width=7.25):
    path = SHOTS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(caption)
    r.italic = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def page_break(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.34)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in (
    ("Title", 28, DARK_ROSE, 0, 12),
    ("Subtitle", 13, MUTED, 0, 10),
    ("Heading 1", 16, DARK_ROSE, 14, 7),
    ("Heading 2", 11.5, ROSE, 10, 5),
    ("Heading 3", 10.5, INK, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

styles["List Bullet"].paragraph_format.left_indent = Inches(0.22)
styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.14)
styles["List Bullet"].paragraph_format.space_after = Pt(3)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("ASTROROSHNI  /  PRODUCTION SCRIPT  /  REVIEW DRAFT")
hr.bold = True
hr.font.size = Pt(7.5)
hr.font.color.rgb = RGBColor.from_string(ROSE)
add_page_number(section.footer.paragraphs[0])

# Cover — workshop_agenda pattern.
accent = doc.add_table(rows=1, cols=1)
accent.autofit = False
cell = accent.cell(0, 0)
set_cell_width(cell, 9360)
shade(cell, ROSE)
set_cell_margins(cell, 28, 0, 28, 0)
cell.paragraphs[0].add_run("")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(28)
p.paragraph_format.space_after = Pt(5)
r = p.add_run("PARASHARI DESK")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string(ROSE)
r.font.letter_spacing = Pt(1.5) if hasattr(r.font, "letter_spacing") else None

p = doc.add_paragraph(style="Title")
p.add_run("Video Explainer\nProduction Script")
p = doc.add_paragraph(style="Subtitle")
p.add_run("Aman voice · synchronized cursor choreography · review draft")

add_callout(
    doc,
    "Production objective",
    "Create a calm, credible guided tour for astrologers: demonstrate how the desk moves from chart context to timing evidence, while every pointer movement, click, zoom and spoken explanation remains synchronized.",
    fill=SAND,
    accent=DARK_ROSE,
)

metrics = doc.add_table(rows=1, cols=4)
metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
metrics.autofit = False
for cell, heading, value in zip(
    metrics.rows[0].cells,
    ("TARGET RUNTIME", "FRAME", "VOICE", "DELIVERABLES"),
    ("8:00", "1920 × 1080", "Aman · original", "MP4 + SRT"),
):
    set_cell_width(cell, 2280)
    shade(cell, IVORY)
    set_cell_margins(cell, 130, 105, 130, 105)
    set_cell_border(cell, top={"val": "single", "sz": "5", "color": LINE}, bottom={"val": "single", "sz": "5", "color": LINE}, left={"val": "single", "sz": "5", "color": LINE}, right={"val": "single", "sz": "5", "color": LINE})
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p0.add_run(heading + "\n")
    rr.bold = True
    rr.font.size = Pt(7.2)
    rr.font.color.rgb = RGBColor.from_string(ROSE)
    rv = p0.add_run(value)
    rv.bold = True
    rv.font.size = Pt(10.2)
    rv.font.color.rgb = RGBColor.from_string(INK)

doc.add_paragraph()
p = doc.add_paragraph("Approval gates", style="Heading 2")
for item in (
    "Approve the spoken script and the order of features.",
    "Approve terminology, captions and pronunciation pickups.",
    "Record the interface only after the first two gates are signed off.",
):
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(item)

add_callout(doc, "Current status", "Script and choreography only. No final interface recording or final narration mix is created at this stage.", fill="F4F7F4", accent=GREEN)

page_break(doc)

doc.add_paragraph("Production language", style="Heading 1")
doc.add_paragraph("The film should feel like an expert showing another expert a well-organized instrument—not like a rapid feature advertisement.")

standards = [
    ("Cursor", "Move in 450–700 ms with a gentle ease. Pause 250–350 ms over the destination. Show a restrained 350 ms click ripple."),
    ("Synchronization", "The pointer leads the spoken label by about 250 ms. The click lands on the action word. Explain the result only after the state change finishes."),
    ("Zoom", "Use controlled 108–118% push-ins only when text or evidence matters. Return to 100% before changing sections."),
    ("Scroll", "Use one deliberate scroll gesture at a time. Let the content settle. Never scrub or hunt for a control on screen."),
    ("Captions", "Two lines maximum, sentence case, exact product spelling. Keep captions clear of bottom navigation and evidence cards."),
    ("Sound", "Aman voice remains foreground. Optional music must be subtle, instrumental and at least 18 dB below narration."),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for i, h in enumerate(("ELEMENT", "RULE")):
    c = table.rows[0].cells[i]
    set_cell_width(c, (1700, 7420)[i])
    shade(c, ROSE)
    c.paragraphs[0].add_run(h).bold = True
    for rr in c.paragraphs[0].runs:
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor.from_string(WHITE)
for idx, (element, rule) in enumerate(standards):
    row = table.add_row()
    prevent_row_split(row)
    for i, value in enumerate((element, rule)):
        c = row.cells[i]
        set_cell_width(c, (1700, 7420)[i])
        set_cell_margins(c, 90, 120, 90, 120)
        shade(c, IVORY if idx % 2 == 0 else WHITE)
        rr = c.paragraphs[0].add_run(value)
        rr.font.size = Pt(9)
        rr.bold = i == 0
        rr.font.color.rgb = RGBColor.from_string(ROSE if i == 0 else INK)
        set_cell_border(c, top={"val": "single", "sz": "3", "color": LINE}, bottom={"val": "single", "sz": "3", "color": LINE}, left={"val": "single", "sz": "3", "color": LINE}, right={"val": "single", "sz": "3", "color": LINE})

doc.add_paragraph("Recording setup", style="Heading 2")
for item in (
    "Desktop: 1600 × 900 browser viewport captured into a 1920 × 1080 master; browser zoom 100%. Use the prepared test native, not private customer data.",
    "Initial as-of date: 09 August 2026. Keep the same native and as-of context throughout, except where the script deliberately demonstrates date synchronization.",
    "Mobile inserts: 390 × 844 viewport. Record separately and place within a clean device frame; do not simulate mobile by squeezing the desktop capture.",
    "Hide developer tools, bookmarks, browser extensions, notification badges and personal account details.",
):
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(item)

doc.add_paragraph("Terminology and voice handling", style="Heading 2")
add_callout(
    doc,
    "Approved voice rule",
    "Use the original Aman voice and its original tone and pace. Keep visible spellings exact: Parashari, Vimshottari, Shadbala, Ashtakavarga, Nakshatra and Dasha. If a term is mispronounced, make a clean word-level pickup and review it; do not switch voices or insert distorted phonetic spellings.",
)

scenes = [
    {
        "number": 1,
        "title": "Open with the predictive workflow",
        "time": "00:00–00:22",
        "duration": "22 seconds",
        "goal": "Establish the desk as one connected workspace for chart context, timing and evidence.",
        "start": "Branded title card, then desktop overview loaded with the test native.",
        "rows": [
            ("00:00", "Fade up title: ‘Parashari Desk — from chart to timing evidence.’ No cursor.", "Parashari Desk"),
            ("00:05", "Cross-dissolve to the full desk at 100%. Cursor rests outside the workspace.", "brings the chart, dashas, transits and event timing into one connected workspace."),
            ("00:13", "Slow 106% push-in centred on the charts and activation dock.", "This tour shows how an astrologer can move from context to a specific timing window—and inspect the evidence behind it."),
        ],
        "narration": "Parashari Desk brings the chart, dashas, transits and event timing into one connected workspace. This tour shows how an astrologer can move from context to a specific timing window—and inspect the evidence behind it.",
        "notes": ["Keep the opening restrained: no feature montage and no fast cursor movement.", "On-screen subtitle: ‘Chart context · timing windows · inspectable evidence’."],
        "image": "01-desktop-overview.png",
        "caption": "Reference frame — complete desktop workspace.",
    },
    {
        "number": 2,
        "title": "Orient the viewer",
        "time": "00:22–00:55",
        "duration": "33 seconds",
        "goal": "Explain the fixed header, primary modes and desk layout before interacting.",
        "start": "Desktop overview at 100%; Activation mode open.",
        "rows": [
            ("00:00", "Cursor glides across desk title and native identity, without clicking.", "The header keeps the native and working context visible."),
            ("00:08", "Trace across Chart, Dasha, Act and More; pause over Act.", "The main modes separate chart study, dasha browsing, activations and classical tools."),
            ("00:18", "Cursor draws a gentle Z-path: chart row, dasha browser, right evidence dock.", "On desktop, charts stay in view while the lower dasha browser and the right-side timing workspace remain available together."),
        ],
        "narration": "The header keeps the native and working context visible. The main modes separate chart study, dasha browsing, activations and classical tools. On desktop, charts stay in view while the lower dasha browser and the right-side timing workspace remain available together.",
        "notes": ["Use a faint rose focus halo following the cursor path; do not darken the rest of the interface.", "Do not click a main mode yet—the purpose here is spatial orientation."],
    },
    {
        "number": 3,
        "title": "Set the as-of moment",
        "time": "00:55–01:32",
        "duration": "37 seconds",
        "goal": "Demonstrate that a single as-of moment synchronizes time-dependent views.",
        "start": "Header visible; as-of date remains 09 August 2026.",
        "rows": [
            ("00:00", "Cursor moves to the date field and pauses. Add 112% crop on the navigator.", "The as-of navigator is the temporal anchor for the desk."),
            ("00:09", "Move across year, month, week, day and hour controls; do not click each control.", "It can move by year, month, week, day or hour, while Now returns to the present moment."),
            ("00:20", "Click the next-day control once. Hold until charts and timing state finish updating.", "When the date changes, the transit chart and time-dependent activation evidence update together."),
            ("00:30", "Click previous-day once to restore 09 August. Pull back to 100%.", "The natal chart remains fixed, so comparison stays clear."),
        ],
        "narration": "The as-of navigator is the temporal anchor for the desk. It can move by year, month, week, day or hour, while Now returns to the present moment. When the date changes, the transit chart and time-dependent activation evidence update together. The natal chart remains fixed, so comparison stays clear.",
        "notes": ["Wait for all loading states to settle before the final sentence.", "Caption the control names only once; do not duplicate labels already visible in the UI."],
    },
    {
        "number": 4,
        "title": "Read charts and activated houses",
        "time": "01:32–02:20",
        "duration": "48 seconds",
        "goal": "Show the chart workspace, divisional access, insight interaction and activation colouring.",
        "start": "Chart mode; D1 and Transit visible with black chart boundaries.",
        "rows": [
            ("00:00", "Click Chart. Sweep across D1, D9, D10 and Transit selectors.", "Chart mode gives direct access to the natal chart, divisional charts and transit context."),
            ("00:12", "Click D9, hold 1.5 seconds, then click D10. Return to D1.", "Divisional charts open in the same workspace, so comparison does not break the reading flow."),
            ("00:24", "Hover an activated D1 house. Click it once to open the insight dock.", "In D1 and Transit, highlighted houses show which areas are active at the selected moment."),
            ("00:36", "Move to the matching Transit house, then the insight content. No further click.", "A house can be opened for its supporting factors, while the two charts preserve the natal and moving-sky relationship."),
        ],
        "narration": "Chart mode gives direct access to the natal chart, divisional charts and transit context. Divisional charts open in the same workspace, so comparison does not break the reading flow. In D1 and Transit, highlighted houses show which areas are active at the selected moment. A house can be opened for its supporting factors, while the two charts preserve the natal and moving-sky relationship.",
        "notes": ["Use a small, persistent legend for activated-house colours before the first hover.", "Do not imply that colour alone predicts an event; call it an activity signal."],
    },
    {
        "number": 5,
        "title": "Browse dashas without losing context",
        "time": "02:20–03:02",
        "duration": "42 seconds",
        "goal": "Show the four dasha systems, five timing levels and date synchronization.",
        "start": "Dasha mode, Vimshottari selected, five timing columns visible.",
        "rows": [
            ("00:00", "Click Dasha. Cursor pauses over Vim, Yog, Kal and Cha tabs.", "The Dasha Browser places four timing systems in one consistent view."),
            ("00:10", "Trace the five columns from Maha through Prana; no click.", "Within a system, five levels—from Mahadasha to Prana—remain visible at the same time."),
            ("00:20", "Click one future Pratyantar row. Wait for selection and date update.", "Selecting a period moves the as-of moment to that interval and updates the linked chart and activation context."),
            ("00:33", "Use the breadcrumb to move one level back; keep cursor still after click.", "Breadcrumbs make it easy to step outward without losing the chosen sequence."),
        ],
        "narration": "The Dasha Browser places four timing systems in one consistent view. Within a system, five levels—from Mahadasha to Prana—remain visible at the same time. Selecting a period moves the as-of moment to that interval and updates the linked chart and activation context. Breadcrumbs make it easy to step outward without losing the chosen sequence.",
        "notes": ["On-screen labels may abbreviate the systems exactly as the product does; captions should use their full names when introduced.", "Keep the selected row centred after the update."],
    },
    {
        "number": 6,
        "title": "Use the Activation Timeline",
        "time": "03:02–03:50",
        "duration": "48 seconds",
        "goal": "Explain current/next windows, confidence states and the expanded timing detail.",
        "start": "Act mode, Timeline selected, legend collapsed to one line.",
        "rows": [
            ("00:00", "Click Act, then Timeline. Cursor points across Strong, Active, Period and Quiet.", "The Activation Timeline organizes the selected period into readable states."),
            ("00:11", "Click ‘What do these mean?’, pause one second, then collapse it.", "The compact legend expands only when its definitions are needed."),
            ("00:21", "Click the first Next window. The detail panel expands to use the full available height.", "Selecting a window opens its timing detail, activated houses and supporting evidence."),
            ("00:34", "Move over the ‘Set as-of’ action but do not click yet; then return to the window heading.", "This is where a broad period becomes a specific, inspectable timing window."),
        ],
        "narration": "The Activation Timeline organizes the selected period into readable states: Strong, Active, Period and Quiet. The compact legend expands only when its definitions are needed. Selecting a window opens its timing detail, activated houses and supporting evidence. This is where a broad period becomes a specific, inspectable timing window.",
        "notes": ["The expanded panel must visibly fill the dock; avoid any empty lower half.", "Let the viewer read the date range for at least 1.5 seconds."],
        "image": "02-activation-timeline.png",
        "caption": "Reference frame — activation timeline and timing evidence dock.",
    },
    {
        "number": 7,
        "title": "Search for a Job Change window",
        "time": "03:50–04:42",
        "duration": "52 seconds",
        "goal": "Demonstrate event-specific search, year selection, result ranking and Full View.",
        "start": "Act mode; Focus tab available; Job change is the default life event.",
        "rows": [
            ("00:00", "Click Focus. Pause over Life event and Year selectors.", "Focus turns the activation engine into an event-specific search."),
            ("00:10", "Open Life event; show Job change and Health, then select Job change.", "Choose a life event and a year. Job Change evaluates career, transition, dasha and independent timing evidence together."),
            ("00:24", "Select 2026, optionally leave developing windows off, click Find windows.", "Find windows returns qualified periods rather than labelling every activated house as the event."),
            ("00:36", "Hover the leading result; click Full view. Hold on the modal title and result count.", "Full View gives the search room to breathe and keeps the evidence readable."),
        ],
        "narration": "Focus turns the activation engine into an event-specific search. Choose a life event and a year. Job Change evaluates career, transition, dasha and independent timing evidence together. Find windows returns qualified periods rather than labelling every activated house as the event. Full View gives the search room to breathe and keeps the evidence readable.",
        "notes": ["Do not show the Timeline, Map or Double Transit tabs inside Full View.", "Keep the Find windows button compact in the recording; avoid an oversized primary block."],
        "image": "03-focus-event-search.png",
        "caption": "Reference frame — Focus event search.",
    },
    {
        "number": 8,
        "title": "Inspect the full calculation",
        "time": "04:42–05:32",
        "duration": "50 seconds",
        "goal": "Build trust by revealing the scored reasoning in plain language.",
        "start": "Focus Full View open with the leading Job Change result visible.",
        "rows": [
            ("00:00", "Click ‘Show full calculations’. Push in to 114% on the calculation header.", "Every qualified window can reveal a human-readable calculation trace."),
            ("00:10", "Slow vertical scroll through Career anchor, Transition signal and Dasha relevance.", "The trace explains what each test means, whether it passed, and which houses and dasha lords supplied the connection."),
            ("00:25", "Continue through independent timing, outcome support and D10 confirmation.", "Independent transit timing, gains support and D10 confirmation refine the case without hiding the underlying evidence."),
            ("00:40", "Pause on exact contact and scoring summary; return to 108%.", "The score is therefore reviewable: the astrologer can agree, disagree or investigate the factors directly."),
        ],
        "narration": "Every qualified window can reveal a human-readable calculation trace. The trace explains what each test means, whether it passed, and which houses and dasha lords supplied the connection. Independent transit timing, gains support and D10 confirmation refine the case without hiding the underlying evidence. The score is therefore reviewable: the astrologer can agree, disagree or investigate the factors directly.",
        "notes": ["Use slow, constant scrolling—about 180 pixels per second—and pause at every scored section.", "Hide ‘Set as-of’ inside Full View because the underlying desk change is not visible there.", "Never show raw JSON."],
        "image": "04-focus-calculation-trace.png",
        "caption": "Reference frame — readable calculation trace.",
    },
    {
        "number": 9,
        "title": "Demonstrate Health as an expandable life area",
        "time": "05:32–05:58",
        "duration": "26 seconds",
        "goal": "Show that Focus supports another event area without overclaiming medical outcomes.",
        "start": "Close Full View; Focus remains open in the desk.",
        "rows": [
            ("00:00", "Open Life event and choose Health. Keep Year at 2026.", "The same framework also supports Health timing."),
            ("00:09", "Click Find windows. Pause on the explanatory subtitle and leading result.", "Here the result is a period of stronger health-related attention—not a diagnosis or a guaranteed medical event."),
            ("00:19", "Hover ‘Show calculations’ without opening it; no zoom.", "Its evidence remains inspectable in the same format."),
        ],
        "narration": "The same framework also supports Health timing. Here the result is a period of stronger health-related attention—not a diagnosis or a guaranteed medical event. Its evidence remains inspectable in the same format.",
        "notes": ["Display a brief on-screen note: ‘Astrological timing evidence; not medical advice.’", "Do not narrate a health prediction for the test native."],
    },
    {
        "number": 10,
        "title": "Find exact Double Transit periods",
        "time": "05:58–06:40",
        "duration": "42 seconds",
        "goal": "Show range selection, exact windows, filters and synchronization to the desk.",
        "start": "Act mode; Double Transit tab selected.",
        "rows": [
            ("00:00", "Click Double Transit. Cursor pauses over From year and Through year dropdowns.", "Double Transit searches for periods when Jupiter and Saturn activate the same natal house."),
            ("00:11", "Open the year dropdowns briefly, then click Calculate exact windows.", "Choose a range and calculate exact ingress boundaries rather than relying on approximate monthly positions."),
            ("00:23", "Switch Future, filter to a life area, and open one result card.", "Results can be filtered by contact type and life area, with the house, contact mode and interpretation visible on each card."),
            ("00:34", "Click the card’s Set as-of action. Hold on the updated header date.", "Selecting a card moves the desk to that period for chart, transit and dasha inspection."),
        ],
        "narration": "Double Transit searches for periods when Jupiter and Saturn activate the same natal house. Choose a range and calculate exact ingress boundaries rather than relying on approximate monthly positions. Results can be filtered by contact type and life area, with the house, contact mode and interpretation visible on each card. Selecting a card moves the desk to that period for chart, transit and dasha inspection.",
        "notes": ["Let the calculation finish before moving the cursor into the results.", "Do not claim that a double transit guarantees the displayed result; describe it as a timing condition."],
        "image": "05-double-transit.png",
        "caption": "Reference frame — Double Transit browser.",
    },
    {
        "number": 11,
        "title": "Open strength and classical tools",
        "time": "06:40–07:14",
        "duration": "34 seconds",
        "goal": "Introduce Shadbala, Ashtakavarga, dignity and Chara Karakas without a long detour.",
        "start": "More mode available; desk returned to 100% zoom.",
        "rows": [
            ("00:00", "Click More. Move through Dignity, Shadbala, Ashtakavarga and Chara Karakas.", "More brings the supporting classical tools into the same visual language."),
            ("00:10", "Click Shadbala. Pause on strength summary and readable table hierarchy.", "Shadbala presents strength measures in a structured, professional view."),
            ("00:20", "Close it; open Ashtakavarga as a screen. Move across Matrix, SAV, BAV and Predictions, then use Back.", "Ashtakavarga opens as a dedicated screen with Matrix, SAV, BAV and Predictions, plus Birth, Transit and Compare controls."),
        ],
        "narration": "More brings the supporting classical tools into the same visual language. Shadbala presents strength measures in a structured, professional view. Ashtakavarga opens as a dedicated screen with Matrix, SAV, BAV and Predictions, plus Birth, Transit and Compare controls. Dignity and Chara Karakas remain close at hand when the reading needs them.",
        "notes": ["Pronunciation check: Shadbala and Ashtakavarga require a word-level review before final mix.", "Show the Ashtakavarga back button unobstructed below the app header."],
    },
    {
        "number": 12,
        "title": "Show the mobile desk",
        "time": "07:14–07:45",
        "duration": "31 seconds",
        "goal": "Demonstrate that mobile uses focused, continuous bottom tab rows rather than shrinking desktop.",
        "start": "Cut to 390 × 844 mobile capture on Chart mode.",
        "rows": [
            ("00:00", "Show Chart with full-width chart and horizontal divisional tabs at the bottom.", "On mobile, the desk becomes a focused browser rather than a compressed desktop layout."),
            ("00:09", "Swipe the divisional tab row; select D9. Tap Dasha and show its bottom system tabs.", "Horizontally scrollable chart and dasha tabs keep every option reachable without crowding the header."),
            ("00:19", "Tap Act, then More. Show readable cards and bottom sub-tabs; one smooth page scroll.", "Activation and More tools use the same continuous bottom-tab pattern, with readable type and touch-sized controls."),
        ],
        "narration": "On mobile, the desk becomes a focused browser rather than a compressed desktop layout. Horizontally scrollable chart and dasha tabs keep every option reachable without crowding the header. Activation and More tools use the same continuous bottom-tab pattern, with readable type and touch-sized controls.",
        "notes": ["Show a visible finger-drag indicator only during the first horizontal swipe.", "Do not trap the recording inside a nested list scroll; the page must move naturally when dragging on a card."],
        "image": "09-mobile-chart.png",
        "caption": "Reference frame — mobile Chart mode.",
    },
    {
        "number": 13,
        "title": "Close on an auditable reading method",
        "time": "07:45–08:00",
        "duration": "15 seconds",
        "goal": "Summarize the product value without promising deterministic outcomes.",
        "start": "Return to desktop overview; clean state at 100%.",
        "rows": [
            ("00:00", "Three restrained highlights: charts, dasha browser, timing evidence.", "Parashari Desk connects natal context, dasha permission and transit timing."),
            ("00:08", "Fade interface into branded end card. Cursor disappears before fade.", "It helps the astrologer identify meaningful windows—and see why each one qualified."),
        ],
        "narration": "Parashari Desk connects natal context, dasha permission and transit timing. It helps the astrologer identify meaningful windows—and see why each one qualified.",
        "notes": ["End card: ‘Parashari Desk · timing evidence you can inspect’. Keep on screen for two seconds after narration.", "Optional footer: ‘Astrological analysis is interpretive and does not guarantee events.’"],
    },
]

for scene in scenes:
    page_break(doc)
    add_scene_header(doc, scene["number"], scene["title"], scene["time"], scene["duration"], scene["goal"], scene["start"])
    add_cue_table(doc, scene["rows"])
    add_narration(doc, scene["narration"])
    add_editor_notes(doc, scene["notes"])
    if scene.get("image"):
        add_label_value(doc, "Visual reference", f"Appendix · {scene['image']}")

page_break(doc)
doc.add_paragraph("Visual reference board · desktop", style="Heading 1")
doc.add_paragraph("Use these frames for composition and control location. Capture the live interface for the film; do not animate the still images themselves.")

desktop_refs = [
    ("01-desktop-overview.png", "Overall desk · Scenes 1–2"),
    ("02-activation-timeline.png", "Activation Timeline · Scene 6"),
    ("03-focus-event-search.png", "Focus search · Scene 7"),
    ("04-focus-calculation-trace.png", "Calculation trace · Scene 8"),
    ("05-double-transit.png", "Double Transit · Scene 10"),
    ("06-shadbala.png", "Shadbala · Scene 11"),
]
for idx, (filename, caption) in enumerate(desktop_refs):
    if idx == 3:
        doc.add_paragraph("Visual reference board · desktop continued", style="Heading 1")
    path = SHOTS / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(path), width=Inches(4.5))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        rr = cp.add_run(caption)
        rr.bold = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor.from_string(MUTED)

page_break(doc)
doc.add_paragraph("Visual reference board · mobile", style="Heading 1")
doc.add_paragraph("Record the mobile views as native responsive screens. Bottom navigation stays unobstructed and every horizontal tab row must visibly respond to touch.")
mobile_table = doc.add_table(rows=1, cols=3)
mobile_table.alignment = WD_TABLE_ALIGNMENT.CENTER
mobile_table.autofit = False
mobile_refs = [
    ("09-mobile-chart.png", "Chart"),
    ("10-mobile-dasha.png", "Dasha"),
    ("11-mobile-activations.png", "Act"),
]
for cell, (filename, caption) in zip(mobile_table.rows[0].cells, mobile_refs):
    set_cell_width(cell, 3040)
    set_cell_margins(cell, 90, 70, 90, 70)
    shade(cell, IVORY)
    set_cell_border(cell, top={"val": "single", "sz": "4", "color": LINE}, bottom={"val": "single", "sz": "4", "color": LINE}, left={"val": "single", "sz": "4", "color": LINE}, right={"val": "single", "sz": "4", "color": LINE})
    path = SHOTS / filename
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        p.add_run().add_picture(str(path), width=Inches(1.75))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p2.add_run(caption)
    rr.bold = True
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor.from_string(ROSE)

page_break(doc)
doc.add_paragraph("Voice, pronunciation and caption sheet", style="Heading 1")
doc.add_paragraph("This sheet is the required preflight for the approved Aman voice. Display text and captions always use the product spelling; pronunciation fixes are handled as clean pickups.")

terms = [
    ("Parashari Desk", "Product name; use this exact visible spelling.", "Listen for a clear middle ‘sha’ and final ‘ree’."),
    ("Vimshottari", "Dasha system.", "Review once in isolation and once in the Scene 5 sentence."),
    ("Shadbala", "Sixfold planetary strength.", "Must not sound like ‘Shadbolu’. Record a pickup if needed."),
    ("Ashtakavarga", "Point-based transit framework.", "Must not end as ‘vargu’. Record a pickup if needed."),
    ("Nakshatra", "Lunar mansion.", "Keep all consonants clear; do not over-slow the sentence."),
    ("Dasha", "Planetary period.", "Use one consistent pronunciation throughout."),
    ("Mahadasha", "Major dasha period.", "Keep natural rhythm; avoid splitting it into two unnatural words."),
    ("Pratyantar", "Sub-period level.", "Test before recording Scene 5."),
    ("Chara Karakas", "Jaimini significators.", "Review the full phrase as spoken in Scene 11."),
]
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [1900, 2900, 4320]
for i, text in enumerate(("TERM", "ON-SCREEN MEANING", "VOICE REVIEW")):
    c = table.rows[0].cells[i]
    set_cell_width(c, widths[i])
    shade(c, ROSE)
    rr = c.paragraphs[0].add_run(text)
    rr.bold = True
    rr.font.size = Pt(7.8)
    rr.font.color.rgb = RGBColor.from_string(WHITE)
for idx, row_data in enumerate(terms):
    row = table.add_row()
    prevent_row_split(row)
    for i, value in enumerate(row_data):
        c = row.cells[i]
        set_cell_width(c, widths[i])
        set_cell_margins(c, 85, 110, 85, 110)
        shade(c, IVORY if idx % 2 == 0 else WHITE)
        rr = c.paragraphs[0].add_run(value)
        rr.font.size = Pt(8.7)
        rr.bold = i == 0
        rr.font.color.rgb = RGBColor.from_string(ROSE if i == 0 else INK)
        set_cell_border(c, top={"val": "single", "sz": "3", "color": LINE}, bottom={"val": "single", "sz": "3", "color": LINE}, left={"val": "single", "sz": "3", "color": LINE}, right={"val": "single", "sz": "3", "color": LINE})

add_callout(
    doc,
    "Pickup workflow",
    "Generate a short Aman pronunciation strip for the flagged terms. Approve it before the full read. If a term fails inside a sentence, replace only that sentence or word with a matched Aman pickup, then crossfade and loudness-match it. Do not use another voice as a patch.",
    fill=SAND,
    accent=DARK_ROSE,
)

doc.add_paragraph("Caption rules", style="Heading 2")
for item in (
    "Export a separate UTF-8 SRT and also prepare a review copy with burned-in captions.",
    "Keep product terminology exactly as shown in the UI; do not spell captions phonetically.",
    "Maximum two lines and approximately 42 characters per line. Break at phrases, never inside a term.",
    "Use white text on a 78% dark neutral plate; maintain at least 5% safe margin from all frame edges.",
):
    q = doc.add_paragraph(style="List Bullet")
    q.add_run(item)

page_break(doc)
doc.add_paragraph("Final recording and review checklist", style="Heading 1")

check_sections = {
    "Before capture": [
        "Approved Aman pronunciation strip is signed off.",
        "Test native contains no private customer data.",
        "Desktop and mobile viewports match the production setup.",
        "All intended screens load without warnings, skeletons or clipped controls.",
    ],
    "Picture and interaction": [
        "Pointer leads the spoken control label by about 250 ms.",
        "Click lands on the spoken action word within ±150 ms.",
        "Every state change finishes before narration explains its result.",
        "No aimless cursor travel, nested-scroll trap, tooltip obstruction or hidden back button.",
        "Chart boundaries, evidence text and mobile controls remain readable at 1080p.",
    ],
    "Narration and terminology": [
        "Aman voice is the original approved voice, tone and pace.",
        "Parashari, Vimshottari, Shadbala, Ashtakavarga and Nakshatra are reviewed in context.",
        "No clipped breaths, abrupt word replacements or audible changes in room tone.",
        "The script describes timing evidence, not guaranteed events or medical diagnosis.",
    ],
    "Delivery": [
        "Master: H.264 MP4, 1920 × 1080, 30 fps, high-bitrate web delivery.",
        "Audio: AAC 48 kHz; narration integrated around −16 LUFS, true peak below −1 dBTP.",
        "Captions: reviewed SRT plus burned-in review export.",
        "Final watch completed once with sound and once muted for visual clarity.",
    ],
}

for heading, items in check_sections.items():
    doc.add_paragraph(heading, style="Heading 2")
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("☐  " + item)
        run.font.size = Pt(9.4)

add_callout(
    doc,
    "Review decision",
    "Approve or annotate this document first. After approval, record a 20–30 second synchronization proof—one cursor move, one click, one state change and one Sanskrit term—before producing the complete eight-minute film.",
    fill="F4F7F4",
    accent=GREEN,
)

doc.core_properties.title = "Parashari Desk Video Explainer Production Script"
doc.core_properties.subject = "Aman narration, synchronized cursor choreography and review checklist"
doc.core_properties.author = "AstroRoshni"
doc.core_properties.keywords = "Parashari Desk, video, narration, storyboard, Aman, astrology"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
