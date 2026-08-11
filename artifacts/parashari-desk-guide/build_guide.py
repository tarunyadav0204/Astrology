from pathlib import Path
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
DERIVED = ROOT / "derived"
DERIVED.mkdir(exist_ok=True)
OUT = ROOT / "Parashari_Desk_Feature_Guide.docx"

ROSE = "A5123B"
ROSE_DARK = "7F0D2D"
INK = "302725"
MUTED = "766761"
IVORY = "FFF9F6"
PALE_ROSE = "F7E7EC"
SAND = "F2E9E2"
GREEN = "DDEAE2"
LINE = "DDCFC6"
WHITE = "FFFFFF"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in kwargs:
            tag = f"w:{edge}"
            edge_el = borders.find(qn(tag))
            if edge_el is None:
                edge_el = OxmlElement(tag)
                borders.append(edge_el)
            for key, value in kwargs[edge].items():
                edge_el.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def crop(src, name, box):
    dst = DERIVED / name
    with Image.open(SHOTS / src) as image:
        image.crop(box).save(dst, quality=95)
    return dst


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ASTROROSHNI  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Title", 28, ROSE_DARK, 0, 12),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 18, ROSE, 4, 9),
        ("Heading 2", 13, INK, 10, 6),
        ("Heading 3", 11, ROSE_DARK, 8, 4),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = style_name != "Subtitle"
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    cap = styles["Caption"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = rgb(MUTED)
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(7)

    # Word's built-in Subtitle style carries a blue bottom border in some
    # templates; remove it so the guide uses only the AstroRoshni rose rule.
    for style_name in ("Title", "Subtitle"):
        style_p_pr = styles[style_name]._element.get_or_add_pPr()
        style_border = style_p_pr.find(qn("w:pBdr"))
        if style_border is not None:
            style_p_pr.remove(style_border)

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = "PARASHARI DESK  /  FEATURE & WORKFLOW GUIDE"
        hp.style = styles["Caption"]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.bold = True
        hp.runs[0].font.color.rgb = rgb(ROSE)
        add_page_number(section.footer.paragraphs[0])
    return doc


def add_rule(doc, color=ROSE, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(ROSE)
    r.font.letter_spacing = Pt(1)
    return p


def add_label(doc, text, fill=PALE_ROSE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_width(table.cell(0, 0), 9360)
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, 100, 150, 100, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = rgb(ROSE_DARK)
    return table


def add_bullets(doc, items, compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(1 if compact else 5)
        p.paragraph_format.line_spacing = 1.1
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_image(doc, path, width=6.9, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph(caption, style="Caption")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_two_images(doc, left, right, captions, widths=(2.55, 2.55)):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, w in enumerate((4680, 4680)):
        set_cell_width(table.cell(0, idx), w)
        set_cell_width(table.cell(1, idx), w)
    for i, (img, width) in enumerate(zip((left, right), widths)):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell_margins(cell, 20, 60, 20, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Inches(width))
        c = table.cell(1, i)
        cell_margins(c, 20, 80, 20, 80)
        cp = c.paragraphs[0]
        cp.style = "Caption"
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(captions[i])
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"val":"nil"}, bottom={"val":"nil"}, start={"val":"nil"}, end={"val":"nil"})
    return table


def new_page(doc, title, kicker=None, intro=None):
    doc.add_page_break()
    if kicker:
        add_kicker(doc, kicker)
    doc.add_heading(title, level=1)
    add_rule(doc, ROSE, 8)
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(8)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2100, 2700, 4560]
    headers = ["Area", "Feature", "What it gives the astrologer"]
    for i, (w, text) in enumerate(zip(widths, headers)):
        set_cell_width(table.cell(0, i), w)
        cell = table.cell(0, i)
        shade(cell, ROSE_DARK)
        cell_margins(cell)
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = rgb(WHITE)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (w, text) in enumerate(zip(widths, row)):
            set_cell_width(cells[i], w)
            cell_margins(cells[i])
            if row_index % 2 == 1:
                shade(cells[i], IVORY)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(text)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"val":"single","sz":"4","color":LINE},
                bottom={"val":"single","sz":"4","color":LINE},
                start={"val":"single","sz":"4","color":LINE},
                end={"val":"single","sz":"4","color":LINE})
    return table


# Exact crops preserve the product UI while focusing attention on dense controls.
header_crop = crop("01-desktop-overview.png", "desktop-header.png", (0, 0, 1600, 180))
chart_crop = crop("01-desktop-overview.png", "desktop-charts.png", (0, 150, 1125, 640))
dasha_crop = crop("01-desktop-overview.png", "desktop-dashas.png", (0, 610, 1125, 900))

doc = setup_doc()

# Cover — editorial cover pattern with branded hero image.
add_kicker(doc, "AstroRoshni professional reference")
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run("Parashari Desk\nFeature & Workflow Guide")
sp = doc.add_paragraph(style="Subtitle")
sp.add_run("A visual guide to charts, dashas, activation timing, event-focused analysis and classical strength tools—across desktop and mobile web.")
add_rule(doc, ROSE, 18)
add_image(doc, SHOTS / "01-desktop-overview.png", 6.92)
cap = doc.add_paragraph("PRODUCT CAPTURE  •  DESKTOP WORKSPACE  •  09 AUG 2026", style="Caption")
cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
cap.runs[0].font.bold = True
cap.runs[0].font.color.rgb = rgb(ROSE)
add_label(doc, "Built for working astrologers: one as-of date, linked charts and dashas, inspectable timing evidence.")
p = doc.add_paragraph("Document version 1.0  |  Screenshots use a test chart; values change with the selected native and as-of time.")
p.style = "Caption"

new_page(doc, "The desk at a glance", "Orientation", "Parashari Desk keeps natal context, the selected as-of sky, period lords and predictive evidence in one coordinated workspace.")
add_image(doc, SHOTS / "01-desktop-overview.png", 6.92, "Desktop overview: four-chart canvas, five-level dasha browser, analysis dock and activated-house colouring.")
add_bullets(doc, [
    ("One temporal spine. ", "The as-of controls drive transits, activation timing and period-sensitive tools."),
    ("Four working zones. ", "Header intelligence, charts, dasha browser and analysis/activation dock stay linked."),
    ("Progressive disclosure. ", "Dense calculations remain available without overwhelming the primary reading surface."),
    ("Responsive continuity. ", "Mobile reorganises the same work into Chart, Dasha, Act and More tabs."),
])

new_page(doc, "Header intelligence & temporal control", "1 · Control spine", "The header is not decoration: it establishes the native, the working date and the contextual facts used throughout the desk.")
add_image(doc, header_crop, 6.92, "Header detail: as-of navigation, birth Panchanga, special points, planetary conditions and special lagnas.")
add_bullets(doc, [
    ("Native context. ", "Return Home, switch to KP or Nadi desks, or replace the active birth chart without leaving the workspace."),
    ("As-of navigator. ", "Jump by year, month, week, day or hour; return to Now; edit date and time directly."),
    ("Birth Panchanga. ", "Tithi, vāra, nakshatra and pada, yoga and karaṇa remain visible for immediate context."),
    ("Special points. ", "Yogi/Avayogi, Dagdha, Tithi Śūnya, Badhaka, Maraka, Rāśi type, Bhrigu Bindu, Pushkara, Mudakku Nakshatra and Gandanta."),
    ("Conditions & lagnas. ", "Planetary state labels plus AL, UL, A7 and Indu Lagna provide fast interpretive anchors."),
])

new_page(doc, "Charts that stay in context", "2 · Chart workspace", "The desktop canvas supports comparison; the mobile browser exposes every chart without forcing a long vertical stack.")
add_image(doc, chart_crop, 6.25, "D1, D9, D10 and Transit in the desktop reading canvas. Activated houses are shaded in D1 and Transit when activation mode is on.")
add_bullets(doc, [
    ("Complete divisional set. ", "D1, D2, D3, D4, D7, D9, D10, D12, D16, D20, D24, D27, D30, D40, D45 and D60, plus Kārakāṁśa and Swāṁśa."),
    ("Chart controls. ", "North/South chart orientation and full-screen inspection are available from the chart header."),
    ("House insight. ", "Selecting a house opens contextual interpretation in the insight dock."),
    ("Activation colours. ", "Strong, Active and Period states are applied only to D1 and Transit, with a visible legend and expandable explanation."),
    ("Transit-specific navigation. ", "The chart-level date navigator appears only when Transit is selected on mobile."),
], compact=True)

new_page(doc, "Dasha browser: hierarchy without losing the thread", "3 · Timing systems", "The dasha browser presents the active chain and adjacent periods together, making it practical to compare the current lord stack with what opens next.")
add_image(doc, dasha_crop, 6.92, "Desktop dasha browser: Mahā, Antar, Pratyantar, Sūkṣma and Prāṇa periods in parallel columns.")
add_bullets(doc, [
    ("Five visible levels. ", "MD, AD, PD, SD and Prāṇa are readable at the same time with active rows clearly marked."),
    ("Four systems. ", "Vimshottari, Yogini, Kalachakra and Chara are available as first-class tabs."),
    ("Period navigation. ", "Selecting a period drills into the next level; the chain summary preserves orientation."),
    ("Date synchronisation. ", "Period selection can move the desk as-of so charts and activation evidence follow the chosen time."),
])

new_page(doc, "Activation Timeline: from state to timing window", "4 · Predictive activation", "Timeline combines dasha portfolios and transits into house-level states, then exposes the reason each timing slice opens and ends.")
add_image(doc, SHOTS / "02-activation-timeline.png", 6.25, "Activation mode on desktop: D1, D9 and Transit above a timeline with a full-height timing-window detail pane.")
add_bullets(doc, [
    ("State legend. ", "Strong, Active, Period and Quiet remain a compact one-line key; definitions expand only when needed."),
    ("Now and Next. ", "Current activated houses are separated from upcoming windows."),
    ("Timing boundaries. ", "Each slice explains what opened it and the change that closes it."),
    ("Evidence detail. ", "The selected house shows dasha participation, transit contact, support/challenge tone and interpretive context."),
    ("As-of handoff. ", "A timing window can set the whole desk to its start for chart and dasha inspection."),
], compact=True)

new_page(doc, "Focus: event-specific window finding", "5 · Guided prediction", "Focus asks a concrete life-event question for a selected year, scores qualified windows and keeps the event definition reusable for future life areas.")
add_image(doc, SHOTS / "03-focus-event-search.png", 6.92, "Job Change focus search in full view, showing qualified windows and the event contract used to find them.")
add_bullets(doc, [
    ("Current focus areas. ", "Job Change and Health use explicit, inspectable rules rather than generic house activation alone."),
    ("Year search. ", "Choose a year, optionally include developing windows, and run the bounded calculation."),
    ("Qualified windows. ", "Results expose confidence, peak date, period lords and concise reasons."),
    ("Full view. ", "A dedicated reading surface removes unrelated tabs and gives the result list and calculation pane room to breathe."),
    ("Expandable architecture. ", "The event definition, scoring gates and evidence sections can add future life areas without rebuilding the desk."),
])

new_page(doc, "Calculation trace: make the algorithm inspectable", "6 · Trust layer", "Every Focus result can be audited in plain language. The interface explains why a gate passed, which planets carried it and what independently confirmed the timing.")
add_image(doc, SHOTS / "04-focus-calculation-trace.png", 6.25, "Expanded calculation for a qualified job-change window: score gates, houses, dasha connections and transit timing evidence.")
add_bullets(doc, [
    ("Career anchor. ", "Shows whether employment/service or profession/status is opened by the active dasha."),
    ("Transition signal. ", "Shows concurrent initiative, transformation, disruption or release indicators."),
    ("Dasha relevance. ", "Names the MD/AD/PD lords and their natal portfolios."),
    ("Independent timing. ", "Surfaces direct transit hits, Jupiter–Saturn reinforcement, dasha boundaries, returns and repeated contacts."),
    ("Outcome refinement. ", "H2/H11 support and D10 confirmation refine the likely professional result."),
], compact=True)

new_page(doc, "Double Transit Browser", "7 · Long-range timing", "This calculator enumerates exact Jupiter–Saturn periods in which both planets activate the same natal house, without substituting guessed ephemeris values.")
add_image(doc, SHOTS / "05-double-transit.png", 6.92, "Double Transit view: exact-year search, Current/Future/Past segmentation, filters and fully explained house cards.")
add_bullets(doc, [
    ("Exact windows. ", "Swiss Ephemeris ingress boundaries are returned by the server and displayed with device-timezone timestamps."),
    ("Contact modes. ", "Filter Full double transits, aspect-only contacts or both."),
    ("Life-area filters. ", "Narrow results to career, relationships, wealth, children, property, health, education or travel."),
    ("Interpretive card. ", "Each result names the house, planet positions/aspects, natal sign/lord/occupants and the governing interpretation rule."),
    ("Desk synchronisation. ", "Selecting a card moves the desktop as-of to the window start for immediate chart and dasha verification."),
])

new_page(doc, "Classical strength, presented for comparison", "8 · Strength tools", "Shadbala turns six-component planetary strength into a professional comparison surface while preserving the underlying point and Rūpa totals.")
add_image(doc, SHOTS / "06-shadbala.png", 6.92, "Shadbala modal: leading and lowest relative strength, ranked planet cards, Rūpas, point totals and expandable components.")
add_bullets(doc, [
    ("At-a-glance range. ", "Leading and lowest relative strength frame the chart before detailed comparison."),
    ("Ranked profile. ", "All seven planets show Rūpas, total points and an assessment grade."),
    ("Component detail. ", "Expand a planet to inspect positional, directional, temporal, motional, natural and aspectual strength."),
    ("Supporting strips. ", "Planetary dignity and Shadbala summaries remain available in the More workspace."),
])

new_page(doc, "Ashtakavarga & Jaimini tools", "9 · Classical analysis", "Full-screen tools preserve complex matrices and categorical results without compressing them into unreadable popups.")
add_image(doc, SHOTS / "07-ashtakavarga.png", 6.35, "Ashtakavarga screen: Birth/Transit/Compare modes with Matrix, SAV, BAV and Predictions sections.")
add_image(doc, SHOTS / "08-chara-karakas.png", 4.15, "Chara Karakas: ranked Jaimini significators with role descriptions; available from More alongside planetary dignities.")
add_bullets(doc, [
    ("Ashtakavarga modes. ", "Birth, Transit and Compare keep the reference frame explicit."),
    ("Ashtakavarga sections. ", "Matrix, SAV, BAV and Predictions are direct tabs; wide tables support touch and horizontal scrolling on mobile."),
    ("Chara Karakas. ", "Atmakaraka through Darakaraka are shown in degree order with their interpretive roles."),
    ("Dignities. ", "Exaltation, debilitation and other dignity states are available as a compact summary and detailed tool."),
], compact=True)

new_page(doc, "Mobile web: charts and dashas", "10 · Responsive workspace", "On mobile, the desk becomes four continuous bottom-tab workspaces. Controls stay thumb-reachable while the selected content receives the full viewport.")
add_two_images(doc, SHOTS / "09-mobile-chart.png", SHOTS / "10-mobile-dasha.png", ("Chart: full-width D1, activation legend, insight entry point and horizontally scrollable chart browser.", "Dasha: as-of controls plus five simultaneous period levels and bottom system tabs."), widths=(2.35, 2.35))
add_bullets(doc, [
    ("Chart browser. ", "All divisional charts, Transit, Kārakāṁśa and Swāṁśa remain available in a horizontal bottom row."),
    ("Readable hierarchy. ", "Chart meta is open by default but collapsible; active-house explanation is a prominent action."),
    ("Dasha density. ", "Horizontal level rows retain surrounding periods without forcing a long vertical list."),
    ("System tabs. ", "Vimshottari, Yogini, Kalachakra and Chara form a continuous bottom tab bar."),
], compact=True)

new_page(doc, "Mobile web: activations and analysis", "11 · Responsive workspace", "The mobile Act and More workspaces separate predictive timing from reference tables while preserving fast switching and readable type.")
add_two_images(doc, SHOTS / "11-mobile-activations.png", SHOTS / "12-mobile-more.png", ("Act: compact legend, current houses, upcoming windows and Timeline/Focus/Map/Double tabs.", "More: dignity, Shadbala, Ashtakavarga, Chara Karakas and analysis tables with bottom section tabs."), widths=(2.35, 2.35))
add_bullets(doc, [
    ("Activation lenses. ", "Timeline, Focus, Map and Double Transit are fixed as mobile bottom tabs."),
    ("Natural page scroll. ", "Lists participate in the screen scroll instead of trapping touch gestures in nested panes."),
    ("More sections. ", "House, Positions, Yogas, Friends, Lords, Aspects and Meta are accessible from the bottom section row."),
    ("Professional touch targets. ", "Continuous tab rows, readable typography and thicker selection states replace tiny chips."),
], compact=True)

new_page(doc, "Recommended predictive workflow", "12 · Working method", "Use the desk as a sequence of tests. Each stage narrows the question and gives the next stage a concrete date, house or planet to inspect.")
workflow = [
    ("1", "Establish the natal frame", "Read D1, relevant divisional chart, strengths, dignity and the special-point/meta strips."),
    ("2", "Choose a concrete question", "Use Focus for Job Change or Health; use Timeline/Map for open-ended house activation."),
    ("3", "Confirm dasha permission", "Inspect MD/AD/PD natal portfolios and the five-level chain at the candidate date."),
    ("4", "Confirm independent timing", "Check direct transit contact, repeated natal relationships, exact returns, boundaries and Double Transit."),
    ("5", "Refine strength and outcome", "Use D9/D10 or the relevant varga, Shadbala, dignity and Ashtakavarga support."),
    ("6", "Move the desk as-of", "Open the candidate date so charts, dashas and activation evidence describe the same moment."),
]
table = doc.add_table(rows=0, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for num, title, desc in workflow:
    cells = table.add_row().cells
    for c, w in zip(cells, (700, 2700, 5960)):
        set_cell_width(c, w)
        cell_margins(c, 110, 130, 110, 130)
        set_cell_border(c, bottom={"val":"single","sz":"5","color":LINE})
    shade(cells[0], ROSE)
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(num)
    r.bold = True
    r.font.color.rgb = rgb(WHITE)
    r.font.size = Pt(13)
    r = cells[1].paragraphs[0].add_run(title)
    r.bold = True
    r.font.color.rgb = rgb(ROSE_DARK)
    cells[2].paragraphs[0].add_run(desc)
add_label(doc, "Interpretation principle: activation is timing evidence, not a guaranteed event. Read natal context, dasha permission, transit confirmation and divisional/strength support together.", SAND)

new_page(doc, "Feature index", "Appendix", "A compact inventory of the desk’s visible capabilities and the question each one answers.")
rows = [
    ("Header", "As-of navigation", "What sky and dasha period am I reading right now?"),
    ("Header", "Birth Panchanga & meta", "What birth-time qualities and special points frame the chart?"),
    ("Charts", "D1 + complete varga set", "Where is the promise, and in which life domain does it specialise?"),
    ("Charts", "Kārakāṁśa / Swāṁśa", "What Jaimini soul and divisional ascendant contexts apply?"),
    ("Charts", "Transit + activation colour", "Which natal houses are being timed at the selected moment?"),
    ("Dashas", "Vimshottari / Yogini / Kalachakra / Chara", "Which period framework is currently operative?"),
    ("Activations", "Timeline", "What is active now and what changes next?"),
    ("Activations", "Focus", "Are there qualified windows for a named life event in this year?"),
    ("Activations", "Map", "Which houses are open across the selected period?"),
    ("Activations", "Double Transit", "When do Jupiter and Saturn simultaneously contact a natal house?"),
    ("Evidence", "Calculation trace", "Why did the event engine qualify this window?"),
    ("Strength", "Shadbala & dignity", "How capable is each planet of delivering its indications?"),
    ("Strength", "Ashtakavarga", "How much bindu support exists by house, planet and transit frame?"),
    ("Jaimini", "Chara Karakas", "Which planets carry the seven movable significator roles?"),
    ("Analysis", "House/positions/yogas/friends/lords/aspects/meta", "What supporting relationships and chart facts refine the reading?"),
]
add_feature_table(doc, rows)

doc.core_properties.title = "Parashari Desk Feature & Workflow Guide"
doc.core_properties.subject = "AstroRoshni Parashari Desk product guide"
doc.core_properties.author = "AstroRoshni"
doc.core_properties.keywords = "Parashari, astrology, dasha, transit, activation, AstroRoshni"
doc.save(OUT)
print(OUT)
